#!/c/Users/SYED/AppData/Local/Programs/Python/Python38-32/python

from tkinter import *
from tkinter.filedialog import *
import docx
from tkinter import messagebox

def openFile():

    file_path=askopenfilename(filetypes=[("All Word Documents","*.docx"),("Text files","*.txt"),("All files","*.*")])

    if not file_path:

       return

    try:

       txt_editor.delete("1.0",END)

       if file_path.endswith(".txt"):

          with open(file_path,"r") as file:
               contents=file.read()
               txt_editor.insert("1.0",contents)
               file.close()

       else:

          doc=docx.Document(file_path)

          full_text=[]

          for para in doc.paragraphs:
              full_text.append(para.text)
          contents="\n".join(full_text)

          txt_editor.insert("1.0",contents)

       window.title("Simple Text Editor - {}".format(file_path))

    except:

          messagebox.showerror("ERROR","Not a valid file!")



def saveFile():

    file_path=asksaveasfilename(defaultextension=".txt",filetypes=[("Text files","*.txt"),("All Word Documents","*.docx"),("CSV(Comma delimited)","*.csv")])

    if not file_path :

       return

    text=txt_editor.get("1.0",END)

    if file_path.endswith(".docx"):

       doc=docx.Document()
       doc.add_paragraph(text)
       doc.save(file_path)

    else:

       with open (file_path,"w") as file:

          text=txt_editor.get("1.0",END)
          file.write(text)
          file.close()
    window.title("Simple Text Editor - {}".format(file_path))

window=Tk()
window.title("Simple Text Editor")

window.rowconfigure(0,minsize=500,weight=1)
window.columnconfigure(1,minsize=900,weight=1)

buttonFrame=Frame(master=window)
txt_editor=Text(master=window)

openButton=Button(master=buttonFrame,text="Open",command=openFile)
saveButton=Button(master=buttonFrame,text="Save as...",command=saveFile)

openButton.grid(row=0,column=0,sticky="ew",padx=10,pady=10)
saveButton.grid(row=1,column=0,sticky="ew",padx=10)

buttonFrame.grid(row=0,column=0,sticky="ns")
txt_editor.grid(row=0,column=1,sticky="nswe")
txt_editor.focus_set()

window.mainloop()
